from odoo import models, fields, api
import logging
from datetime import datetime
from io import BytesIO
import base64
import os

try:
    from docx import Document
    from docx.shared import Inches
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from PyPDF2 import PdfWriter, PdfReader
    import subprocess
    import tempfile
except ImportError as e:
    logging.getLogger(__name__).warning(f"Required libraries not installed: {e}. PDF generation will not work.")

_logger = logging.getLogger(__name__)


class AccountMoveWarranty(models.Model):
    _inherit = 'account.move'

    def generate_warranty_pdfs(self):
        """
        Generate warranty PDFs by creating Word documents programmatically with required fields.
        
        Business Rules:
        - Create Word documents programmatically (no external template files)
        - Fill customer name, product brand, warranty period
        - Generate one PDF per product with warranty
        - Exclude product with ID 7884
        - Default to "1" month if warranty field is empty
        
        Returns:
            dict: Action dictionary for file download
        """
        try:
            # Get products from invoice lines
            products = []
            for line in self.invoice_line_ids:
                if line.product_id and line.product_id.id != 7884:
                    products.append(line.product_id)
            
            if not products:
                return {
                    'type': 'ir.actions.client',
                    'tag': 'display_notification',
                    'params': {
                        'title': 'No Products',
                        'message': 'No valid products found for warranty generation.',
                        'type': 'warning',
                    }
                }
            
            # Generate filled warranty PDFs
            pdf_content = self._generate_filled_warranty_pdf(products)
            
            if not pdf_content:
                return {
                    'type': 'ir.actions.client',
                    'tag': 'display_notification',
                    'params': {
                        'title': 'Error',
                        'message': 'Failed to generate warranty PDF. Please check template file and dependencies.',
                        'type': 'danger',
                    }
                }
            
            # Create attachment for download
            filename = f"garancia_{self.name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            attachment = self.env['ir.attachment'].create({
                'name': filename,
                'type': 'binary',
                'datas': base64.b64encode(pdf_content),
                'res_model': 'account.move',
                'res_id': self.id,
                'mimetype': 'application/pdf'
            })
            
            return {
                'type': 'ir.actions.act_url',
                'url': f'/web/content/{attachment.id}?download=true',
                'target': 'self',
            }
            
        except Exception as e:
            _logger.error(f'Error generating warranty PDF for invoice {self.id}: {str(e)}')
            return {
                'type': 'ir.actions.client',
                'tag': 'display_notification',
                'params': {
                    'title': 'Error',
                    'message': f'Error generating warranty PDF: {str(e)}',
                    'type': 'danger',
                }
            }

    def _generate_filled_warranty_pdf(self, products):
        """
        Generate warranty PDF by creating Word documents with product data.
        
        Args:
            products: List of product.product records
            
        Returns:
            bytes: PDF content as bytes
        """
        try:
            # Create a new PDF writer for the final output
            output_pdf = PdfWriter()
            
            # Process each product
            for product in products:
                filled_pdf = self._create_warranty_document(product)
                if filled_pdf:
                    # Add the filled page to output
                    reader = PdfReader(BytesIO(filled_pdf))
                    for page in reader.pages:
                        output_pdf.add_page(page)
            
            # Write final PDF to buffer
            output_buffer = BytesIO()
            output_pdf.write(output_buffer)
            pdf_content = output_buffer.getvalue()
            output_buffer.close()
            
            return pdf_content
            
        except Exception as e:
            _logger.error(f'Error generating filled warranty PDF: {str(e)}')
            return None

    def _create_warranty_document(self, product):
        """
        Create warranty document with product-specific data.
        
        Args:
            product: Product record
            
        Returns:
            bytes: Filled PDF content
        """
        try:
            # Create a new Word document
            doc = Document()
            
            # Get customer name from invoice partner
            customer_name = self.partner_id.name or "___________________"
            
            # Get product brand/name
            product_name = product.name or "________________"
            
            # Get warranty period (default to 1 if empty)
            warranty_period = str(getattr(product, 'x_studio_warranty', '1')).replace(' muaj', '').replace(' month', '').strip()
            if not warranty_period or warranty_period == '':
                warranty_period = '1'
            
            # Get invoice information
            invoice_number = self.name or ''
            invoice_date = self.invoice_date.strftime('%d/%m/%Y') if self.invoice_date else ''
            
            # Create warranty certificate content
            self._add_warranty_content(doc, customer_name, product_name, warranty_period, invoice_number, invoice_date)
            
            # Save document to temporary file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
                doc.save(temp_docx.name)
                temp_docx_path = temp_docx.name
            
            # Convert Word document to PDF
            pdf_content = self._convert_docx_to_pdf(temp_docx_path)
            
            # Clean up temporary file
            os.unlink(temp_docx_path)
            
            return pdf_content
            
        except Exception as e:
            _logger.error(f'Error creating warranty document for product {product.id}: {str(e)}')
            return None

    def _add_warranty_content(self, doc, customer_name, product_name, warranty_period, invoice_number, invoice_date):
        """
        Add warranty certificate content to the document.
        
        Args:
            doc: Word document object
            customer_name (str): Customer name
            product_name (str): Product name
            warranty_period (str): Warranty period in months
            invoice_number (str): Invoice number
            invoice_date (str): Invoice date
        """
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        
        # Title
        title = doc.add_heading('GARANCIA', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add some space
        doc.add_paragraph()
        
        # Customer information section
        customer_para = doc.add_paragraph()
        customer_para.add_run('Emer Mbiemer: ').bold = True
        customer_para.add_run(customer_name)
        
        # Product information section
        product_para = doc.add_paragraph()
        product_para.add_run('Marka: ').bold = True
        product_para.add_run(product_name)
        
        # Warranty period section
        warranty_para = doc.add_paragraph()
        warranty_para.add_run('Afati Garancise: ').bold = True
        warranty_para.add_run(f'{warranty_period} Muaj')
        
        # Invoice information section
        invoice_para = doc.add_paragraph()
        invoice_para.add_run('Numri i Fatures: ').bold = True
        invoice_para.add_run(invoice_number)
        
        date_para = doc.add_paragraph()
        date_para.add_run('Data e Fatures: ').bold = True
        date_para.add_run(invoice_date)
        
        # Add some space
        doc.add_paragraph()
        
        # Warranty terms section
        terms_heading = doc.add_heading('Kushtet e Garancise', level=2)
        
        terms_text = """
        Kjo garancia mbulon defekte ne material dhe punim te produktit.
        Garancia nuk mbulon:
        - Demet e shkaktuara nga perdorimi gabim
        - Demet e shkaktuara nga aksidentet
        - Demet e shkaktuara nga modifikimet e produktit
        - Konsumimin normal te produktit
        
        Per te aktivizuar garancine, kontaktoni:
        Telefon: [NUMRI I TELEFONIT]
        Email: [EMAIL ADRESA]
        Adresa: [ADRESA E PLOTE]
        """
        
        terms_para = doc.add_paragraph(terms_text.strip())
        
        # Add some space
        doc.add_paragraph()
        
        # Signature section
        signature_para = doc.add_paragraph()
        signature_para.add_run('Nenshkrimi: _________________________')
        signature_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        date_signature_para = doc.add_paragraph()
        date_signature_para.add_run('Data: _________________________')
        date_signature_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _convert_docx_to_pdf(self, docx_path):
        """
        Convert Word document to PDF using LibreOffice.
        
        Args:
            docx_path (str): Path to the Word document
            
        Returns:
            bytes: PDF content as bytes
        """
        try:
            # Create temporary PDF file
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
                temp_pdf_path = temp_pdf.name
            
            # Use LibreOffice to convert DOCX to PDF
            cmd = [
                'libreoffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', os.path.dirname(temp_pdf_path),
                docx_path
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            
            if result.returncode != 0:
                _logger.error(f'LibreOffice conversion failed: {result.stderr}')
                return None
            
            # Read the generated PDF
            pdf_path = os.path.splitext(temp_pdf_path)[0] + '.pdf'
            if os.path.exists(pdf_path):
                with open(pdf_path, 'rb') as pdf_file:
                    pdf_content = pdf_file.read()
                
                # Clean up temporary files
                os.unlink(temp_pdf_path)
                if pdf_path != temp_pdf_path:
                    os.unlink(pdf_path)
                
                return pdf_content
            else:
                _logger.error(f'PDF file not generated: {pdf_path}')
                return None
                
        except subprocess.TimeoutExpired:
            _logger.error('LibreOffice conversion timed out')
            return None
        except Exception as e:
            _logger.error(f'Error converting DOCX to PDF: {str(e)}')
            return None


class WarrantyPdfSettings(models.TransientModel):
    _name = 'warranty.pdf.settings'
    _description = 'Warranty PDF Generation Settings'

    exclude_product_ids = fields.Many2many(
        'product.product',
        string='Exclude Products',
        help='Products to exclude from warranty generation'
    )
    
    default_warranty_period = fields.Char(
        string='Default Warranty Period',
        default='1',
        help='Default warranty period in months when product warranty is not set'
    )
    
    template_file = fields.Binary(
        string='Template File',
        help='Upload custom warranty template Word document'
    )
    
    template_filename = fields.Char(
        string='Template Filename'
    )

    @api.model
    def get_settings(self):
        """Get warranty PDF generation settings."""
        config = self.env['ir.config_parameter'].sudo()
        return {
            'exclude_product_ids': config.get_param('warranty_pdf.exclude_product_ids', '7884').split(','),
            'default_warranty_period': config.get_param('warranty_pdf.default_warranty_period', '1'),
        }

    def save_settings(self):
        """Save warranty PDF generation settings."""
        config = self.env['ir.config_parameter'].sudo()
        
        exclude_ids = ','.join(str(pid) for pid in self.exclude_product_ids.ids) if self.exclude_product_ids else '7884'
        config.set_param('warranty_pdf.exclude_product_ids', exclude_ids)
        config.set_param('warranty_pdf.default_warranty_period', self.default_warranty_period or '1')
        
        return {
            'type': 'ir.actions.client',
            'tag': 'display_notification',
            'params': {
                'title': 'Settings Saved',
                'message': 'Warranty PDF settings have been saved successfully.',
                'type': 'success',
            }
        }